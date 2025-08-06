import os, base64, subprocess, tempfile
from pypdf import PdfReader

import gradio as gr
import faiss
from sentence_transformers import SentenceTransformer

from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.llms import Ollama
from langchain.chains import RetrievalQA

from docx import Document
import asyncio

# ---------------------
# 1) CONFIG
# ---------------------
OLLAMA_MODEL = "llama3:8b-instruct-q5_K_M"
EMBED_MODEL  = SentenceTransformer("all-MiniLM-L6-v2")

# ---------------------
# 2) GLOBAL STATE
# ---------------------
DOCUMENT_PAGES = []
CURRENT_PAGE   = 0
QA_CHAIN       = None
RETRIEVER      = None

# ---------------------
# 3) UTILITIES
# ---------------------
def extract_pages(path):
    reader = PdfReader(path)
    return [p.extract_text() or "" for p in reader.pages]

def build_rag(pdf_path):
    global DOCUMENT_PAGES, CURRENT_PAGE, QA_CHAIN, RETRIEVER
    # 1) raw pages
    DOCUMENT_PAGES = extract_pages(pdf_path)
    CURRENT_PAGE   = 0
    # 2) load & chunk
    # Tune chunk_size and chunk_overlap for latency/recall tradeoff. For large docs, try chunk_size=1200, chunk_overlap=100.
    # For even faster search, use faiss-gpu if available.
    loader = PyPDFLoader(pdf_path)
    docs   = loader.load()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800, chunk_overlap=200
    )
    chunks = splitter.split_documents(docs)
    # 3) embeddings & FAISS
    embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    store = FAISS.from_documents(chunks, embedding_model)
    # 4) retriever with MMR
    RETRIEVER = store.as_retriever(
        search_type="mmr", search_kwargs={"k":4, "fetch_k":8}
    )
    # 5) LLM + map-reduce QA
    # To use a GPU-accelerated backend, replace Ollama here with e.g. llama.cpp or vLLM integration.
    llm = Ollama(model=OLLAMA_MODEL, temperature=0.2)
    QA_CHAIN = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="map_reduce",
        retriever=RETRIEVER,
        return_source_documents=False
    )

# ---------------------
# 4) HANDLERS
# ---------------------
async def upload_and_process(pdf_file):
    if pdf_file is None:
        return "❌ No file", "", ""
    try:
        await asyncio.to_thread(build_rag, pdf_file.name)
        # preview first snippet
        snippet = DOCUMENT_PAGES[0][:500] + "…" if DOCUMENT_PAGES else ""
        # embed PDF for preview
        raw = open(pdf_file.name, "rb").read()
        b64 = base64.b64encode(raw).decode()
        embed = (
            f'<object data="data:application/pdf;base64,{b64}" '
            'type="application/pdf" width="100%" height="600px">'
            "PDF preview not available</object>"
        )
        return "✅ Loaded & RAG built", embed, snippet
    except Exception as e:
        return f"❌ Error: {e}", "", ""

async def show_prev_page():
    global CURRENT_PAGE
    if not DOCUMENT_PAGES:
        return "No pages loaded"
    CURRENT_PAGE = max(0, CURRENT_PAGE - 1)
    txt = DOCUMENT_PAGES[CURRENT_PAGE]
    return txt[:500] + "…" if len(txt) > 500 else txt

async def show_next_page():
    global CURRENT_PAGE
    if not DOCUMENT_PAGES:
        return "No pages loaded"
    CURRENT_PAGE = min(len(DOCUMENT_PAGES) - 1, CURRENT_PAGE + 1)
    txt = DOCUMENT_PAGES[CURRENT_PAGE]
    return txt[:500] + "…" if len(txt) > 500 else txt

async def chat_and_retrieve(mode, query, history):
    if QA_CHAIN is None:
        return history, history, "❌ Upload a PDF first"

    history = history or []
    if mode == "Summarize":
        # just summarize full doc
        full_text = "\n\n".join(DOCUMENT_PAGES)
        llm = Ollama(model=OLLAMA_MODEL, temperature=0.2)
        summary = await asyncio.to_thread(llm.predict, f"Summarize in 3 sentences:\n\n{full_text}")
        history.append((f"[Summarize]", summary.strip()))
        return history, history, ""

    # Ask mode
    # 1) retrieve passages
    docs = await asyncio.to_thread(RETRIEVER.get_relevant_documents, query)
    retrieved = "\n\n---\n\n".join(d.page_content for d in docs)

    # 2) answer
    try:
        ans = (await asyncio.to_thread(QA_CHAIN, {"query": query}))["result"].strip()
    except Exception as e:
        ans = f"❌ LLM error: {e}"

    history.append((query, ans))
    return history, history, retrieved

async def export_history(history, filetype="txt"):
    if not history:
        return None
    if filetype == "docx":
        path = "chat_log.docx"
        doc = Document()
        for q, a in history:
            doc.add_paragraph(f"Q: {q}")
            doc.add_paragraph(f"A: {a}")
            doc.add_paragraph("")
        await asyncio.to_thread(doc.save, path)
        return path
    else:
        path = "chat_log.txt"
        with open(path, "w") as f:
            for q, a in history:
                f.write(f"Q: {q}\nA: {a}\n\n")
        return path

# ---------------------
# 5) GRADIO UI
# ---------------------
with gr.Blocks(title="DoChat+ w/ LLaMA 3 RAG") as demo:
    gr.Markdown("## 📄 DoChat+ — PDF Q&A with LLaMA 3 8B (MMR + Map-Reduce)")

    with gr.Row():
        pdf_in     = gr.File(label="Upload PDF", file_types=[".pdf"])
        upload_btn = gr.Button("Upload")
        status     = gr.Textbox(label="Status")

    preview_pdf  = gr.HTML()
    preview_text = gr.Textbox(label="First Snippet", lines=4)

    with gr.Row():
        prev_btn = gr.Button("⯇ Prev")
        next_btn = gr.Button("Next ⯈")

    chatbot       = gr.Chatbot()
    retrieved_ctx = gr.Textbox(label="🔍 Retrieved Passages", lines=6)

    with gr.Row():
        mode     = gr.Radio(["Ask","Summarize"], value="Ask", label="Mode")
        question = gr.Textbox(placeholder="Type here…", label="Your Input")
        send_btn = gr.Button("Send")

    filetype_radio = gr.Radio(["txt", "docx"], value="txt", label="File Type")
    export_btn    = gr.Button("📥 Download Chat Log")
    download_file = gr.File()

    upload_btn.click(upload_and_process, pdf_in, [status, preview_pdf, preview_text])
    prev_btn.click(show_prev_page, [], preview_text)
    next_btn.click(show_next_page, [], preview_text)
    send_btn.click(chat_and_retrieve, [mode, question, chatbot], [chatbot, chatbot, retrieved_ctx])
    export_btn.click(export_history, [chatbot, filetype_radio], download_file)

if __name__ == "__main__":
    demo.launch()
